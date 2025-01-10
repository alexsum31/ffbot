import streamlit as st
import pandas as pd
from re import M, X
from docx import Document
from docx.shared import Mm,Pt,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import datetime
from pathlib import Path
from docx.enum.text import WD_LINE_SPACING
import re
from io import StringIO,BytesIO
import xlsxwriter

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                    
def makefontBig(tableName,col1,row1,sizeA):
    hdr_cells = tableName.rows[row1].cells         
    paragraph =hdr_cells[col1].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size= Pt(sizeA)

def set_col_widths(table):
    widths = (Mm(70), Mm(89),Mm(15))
    for rowD in table.rows:
        for idx, width in enumerate(widths):
            rowD.cells[idx].width = width
            

def makeTable(doc,orderIDnumber,counta,df):
    tempdf=df[df.order_id.isin([orderIDnumber])]
    tempdf.reset_index(inplace=True) 
    aList = []
    QTYList=[]
    index=0
    for index, row in tempdf.iterrows():
        if index==0:
            Tlst = [row['order_id'], 
                'Store: '+row['store'],
                'Channel: '+row['channel'],
                'Sales Order Date: '+ row['order_date'].strftime('%Y-%m-%d')]#, %H:%M:%S'
        aList.append(str(row['ff_modelcode'])+'('+str(row['barcode1'])+')'+row['ff_modeldes'])
        QTYList.append(str(row['qty']))
    d = {'Order info':Tlst,'Details':aList,'Qty':QTYList}
    FinalDetails = pd.DataFrame.from_dict(d, orient='index')
    FinalDetails = FinalDetails.transpose()

    table = doc.add_table(FinalDetails.shape[0]+1, FinalDetails.shape[1])
    set_col_widths(table)
    
    for j in range(FinalDetails.shape[-1]):
          table.cell(0,j).text = FinalDetails.columns[j]

    for i in range(FinalDetails.shape[0]):
         for j in range(FinalDetails.shape[-1]):   
                table.cell(i+1,j).text=str(FinalDetails.values[i,j])
                if str(FinalDetails.values[i,j])=='None':
                     table.cell(i+1,j).text=''
                else:
                    table.cell(i+1,j).text=str(FinalDetails.values[i,j])
                    if  j==1:
                        makefontBig(table,j,i+1,10)
                    if  j ==max(range(FinalDetails.shape[-1])):
                        makefontBig(table,j,i+1,18)
                if i+1==max(range(FinalDetails.shape[0]))+1:
                    set_cell_border(
                    table.cell(i+1,j),
                    bottom={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"}
                    )
     
    makefontBig(table,0,1,18)
    makefontBig(table,0,2,18)
    for row in table.rows:
            row.height = Mm(12.5)
            table.rows[0].height=Mm(5.04)
    # #table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
    a = table.cell(0, 0)
    b = table.cell(4, 0)
    a.merge(b)
    
    mod =1
    mod = counta % 4
    if mod==0:
        doc.add_page_break()
        
def eachtablewriter(T,xlsdataframe,doc):
    List_orderID=xlsdataframe['order_id'].unique()
    for x in List_orderID:
        makeTable(doc,x,T,xlsdataframe)
        T+=1
    
def writedoc(abc):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'DengXian'
    font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(5)
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(15.7)
    section.right_margin = Mm(15.7)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    return doc


        
def get_docx_download_link(docx):
    buffered = BytesIO()
    docx.save(buffered)
    return buffered.getvalue()

def app():
    savedatestr=datetime.date.today()
    T_today=savedatestr.strftime('%Y%m%d')
    st.header('Golden Cosmetices Package Label',divider="rainbow")
    xls_file = st.file_uploader("Upload excel file",type=['.xlsx'])
    if xls_file:
        try:
            df = pd.read_excel(xls_file,dtype={'Content.Store':str}) 
            df.columns= ['channel',
            'order_id',
            'order_date',
            'sr_modelcode',
            'store',
            'barcode1',
            'barcode2',
            'ff_modelcode',
            'ff_modeldes',
            'qty',
            'Sent_Golden_Cosmetic'
                ]

            doc=writedoc("start")
            eachtablewriter(1,df,doc)
            dummy_name=xls_file.name[:len(xls_file.name)-5]+'_'
            ##edited_doc1=doc.save(path+dummy_name+T_today+'.docx')

            st.download_button(label='Download edited document',
                                data= get_docx_download_link(doc),
                                file_name=dummy_name+T_today+'.docx'
                                #mime="application/octet-stream"
                                )
            
        except:
            st.error('Upload file data error',  icon="⚠️")   
           


app()